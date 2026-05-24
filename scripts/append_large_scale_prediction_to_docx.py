#!/usr/bin/env python3
"""Append large-scale runtime predictions to the permutation DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


TARGET = Path(
    "/Users/piotrlunkiewicz/Desktop/Doktorat UAM/Zadania/Permutacja sprawdzamy/"
    "wyniki_RSA_permutacja_vs_bez_2026-05-24/"
    "opracowanie_permutacja_w_RSA_wyjasnienie_i_dyskusja.docx"
)


def style_run(run, *, bold: bool = False, color: str | None = None, size: int = 11) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def main() -> None:
    doc = Document(TARGET)
    title = "Przewidywanie dla duzej symulacji: L=1000 i 10^4 przebiegow"
    existing = "\n".join(p.text for p in doc.paragraphs)
    if title in existing:
        raise SystemExit("Section already exists; not appending duplicate content.")

    doc.add_page_break()
    doc.add_heading(title, level=1)

    p = doc.add_paragraph()
    r = p.add_run("To jest przewidywanie/ekstrapolacja, a nie wynik pelnej symulacji.")
    style_run(r, bold=True, color="1F4D78")

    doc.add_paragraph(
        "Dla L=1000 liczba punktow kratowych wynosi 1 000 000, a liczba mozliwych polozen dimeru przy hard boundary condition wynosi "
        "2*L*(L-1)=1 998 000. To jest okolo 100.9 razy wiecej kandydatow niz dla L=100, gdzie bylo 19 800 kandydatow."
    )
    doc.add_paragraph(
        "Jesli zachowamy te sama logike programu i zalozymy prawie liniowe skalowanie z liczba kandydatow, to pojedynczy przebieg dla L=1000 moglby trwac rzedu 1 sekundy w zoptymalizowanej wersji pythonowej. "
        "Dla serii 10^4 przebiegow oznacza to rzad wielkosci kilku godzin na metode, a dla dwoch metod razem kilka dodatkowych godzin. "
        "To oszacowanie moze sie zmienic w zaleznosci od komputera, implementacji, pamieci RAM i sposobu przechowywania list kandydatow."
    )

    doc.add_heading("Przewidywany czas dla wersji zoptymalizowanych", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Metoda", "Czas na przebieg L=100", "Prognoza L=1000", "10^4 przebiegow", "Komentarz"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
    rows = [
        (
            "bez permutacji, zoptymalizowana",
            "0.01068 s",
            "ok. 1.08 s",
            "ok. 3.0 h",
            "usuwa niemozliwe kandydaty, wiec nie traci duzo czasu na puste proby",
        ),
        (
            "z permutacja",
            "0.00984 s",
            "ok. 0.99 s",
            "ok. 2.7 h",
            "ok. 8% szybciej w pomiarze dla tej implementacji",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value

    doc.add_paragraph(
        "W takim wariancie roznica czasowa miedzy dwiema dobrze zoptymalizowanymi metodami nadal moglaby byc umiarkowana: rzedu kilku do kilkunastu procent. "
        "Dla 10^4 przebiegow roznica 8% oznaczalaby jednak juz nie milisekundy, lecz kilkanascie minut lub wiecej zaoszczedzonego czasu."
    )

    doc.add_heading("Porownanie z naiwnym RSA z odrzuceniami", level=2)
    doc.add_paragraph(
        "Najwieksza roznica pojawilaby sie przy porownaniu permutacji z klasycznym, naiwnym RSA: losuj pozycje i orientacje, sprawdz, odrzuc, losuj ponownie. "
        "Dla L=1000 i blisko jammingu prawdopodobienstwo akceptacji jest bardzo male, wiec ogromna liczba prob nie zmienialaby juz konfiguracji."
    )
    doc.add_paragraph(
        "W takim naiwnym wariancie czas moglby wzrosnac nie o 8%, ale wielokrotnie: potencjalnie od kilku razy do rzedow wielkosci, zwlaszcza jesli kryterium zatrzymania opiera sie na dlugiej serii porazek. "
        "Permutacja ma wtedy zasadnicza przewage: kazdy kandydat jest sprawdzany najwyzej raz, a koniec listy daje jednoznaczny stan jammed."
    )

    doc.add_heading("Przewidywana roznica wynikow fizycznych", level=2)
    doc.add_paragraph(
        "Dla L=1000 i 10^4 przebiegow nie oczekiwalbym systematycznej roznicy miedzy Cp i Cj dla wersji z permutacja i bez permutacji, jesli obie metody losuja te sama przestrzen kandydatow i stosuja to samo kryterium akceptacji. "
        "Wieksza siatka oraz wieksza liczba przebiegow powinny raczej zmniejszyc blad sredniej i jeszcze mocniej pokazac, ze roznice sa losowym szumem Monte Carlo, a nie efektem permutacji."
    )
    doc.add_paragraph(
        "Przy L=1000 rozklad Cj powinien byc wezszy niz dla L=100, bo uklad jest wiekszy i fluktuacje wzgledne maleja. "
        "Przy 10^4 przebiegow SEM spada dodatkowo jak 1/sqrt(R), czyli w porownaniu z 500 przebiegami okolo sqrt(20), mniej wiecej 4.5 raza. "
        "Dlatego nawet bardzo mala roznica moglaby byc liczbowo mierzona dokladniej, ale nadal nie powinna miec znaku systematycznego, jesli algorytmy sa rownowazne."
    )

    doc.add_heading("Podsumowanie w tym samym stylu", level=2)
    doc.add_paragraph(
        "Odpowiedz: tak, z permutacja symulacja bylaby zwykle szybsza rowniez dla bardzo duzego przypadku L=1000 i serii 10^4 przebiegow. "
        "W porownaniu dwoch zoptymalizowanych wersji roznica moglaby pozostac umiarkowana, na przyklad okolo kilku-kilkunastu procent. "
        "W porownaniu z naiwnym RSA z odrzuceniami przewaga permutacji moglaby byc natomiast bardzo duza, bo blisko jammingu klasyczny algorytm wykonuje ogromna liczbe nieudanych prob."
    )
    doc.add_paragraph(
        "Najwazniejszy wniosek pozostaje taki sam: permutacja przyspiesza i porzadkuje obliczenia, ale sama w sobie nie powinna zmieniac wartosci Cp ani Cj, jesli definicja kandydatow, hard boundary condition i reguly akceptacji sa identyczne."
    )

    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
