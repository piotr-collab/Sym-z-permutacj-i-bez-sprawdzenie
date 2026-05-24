#!/usr/bin/env python3
"""Append a speed discussion section to the permutation DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


TARGET = Path(
    "/Users/piotrlunkiewicz/Desktop/Doktorat UAM/Zadania/Permutacja sprawdzamy/"
    "wyniki_RSA_permutacja_vs_bez_2026-05-24/"
    "opracowanie_permutacja_w_RSA_wyjasnienie_i_dyskusja.docx"
)


def set_run_style(run, *, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def main() -> None:
    doc = Document(TARGET)

    existing = "\n".join(p.text for p in doc.paragraphs)
    section_title = "Czy z permutacja symulacja jest szybsza?"
    if section_title in existing:
        raise SystemExit("Section already exists; not appending duplicate content.")

    doc.add_page_break()
    doc.add_heading(section_title, level=1)

    p = doc.add_paragraph()
    r = p.add_run("Krotka odpowiedz: tak, zwykle jest szybciej.")
    set_run_style(r, bold=True, color="1F4D78")

    doc.add_paragraph(
        "Wersja z permutacja tworzy raz losowa kolejnosc wszystkich mozliwych polozen dimeru i przechodzi po tej liscie. "
        "Dzieki temu unika wielokrotnego losowania tych samych nieudanych prob, ktore w klasycznym RSA pojawiaja sie bardzo czesto blisko stanu jammed."
    )
    doc.add_paragraph(
        "Najwiekszy zysk czasowy widac przy porownaniu z naiwnym algorytmem RSA z odrzuceniami: losuj pozycje, losuj orientacje, sprawdz, odrzuc, losuj ponownie. "
        "Pod koniec symulacji prawdopodobienstwo trafienia legalnego miejsca jest male, wiec komputer wykonuje duzo prob, ktore nie zmieniaja konfiguracji."
    )
    doc.add_paragraph(
        "W tej pracy porownanie wykonano jednak z wersja bez permutacji, ktora rowniez byla zoptymalizowana: zamiast losowac nieskonczenie wiele odrzuconych prob, wybierala kolejny dimer z aktualnie legalnych kandydatow i usuwala kandydaty niemozliwe po kazdej akceptacji. "
        "Dlatego roznica czasu jest mniejsza niz w porownaniu z naiwnym RSA."
    )

    doc.add_heading("Pomiar czasu w tej implementacji", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Metoda", "Liczba przebiegow", "Czas laczny", "Czas na przebieg"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
    rows = [
        ("bez permutacji", "100", "1.068 s", "0.01068 s"),
        ("z permutacja", "100", "0.984 s", "0.00984 s"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value

    doc.add_paragraph(
        "W tym konkretnym pomiarze wersja z permutacja byla szybsza o okolo 8%. "
        "Nie nalezy jednak traktowac tej liczby jako uniwersalnej stalej: zalezy ona od implementacji, jezyka programu, sposobu aktualizacji listy kandydatow oraz rozmiaru siatki."
    )
    doc.add_paragraph(
        "Wniosek praktyczny jest nastepujacy: permutacja nie zmienia oczekiwanego wyniku geometrycznego, czyli rozkladu Cp i Cj, ale zwykle upraszcza zatrzymanie symulacji i zmniejsza koszt obliczeniowy, szczegolnie wzgledem klasycznego losowania z odrzuceniami."
    )

    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
