#!/usr/bin/env python3
"""Build a standalone DOCX explaining permutation in the RSA simulation."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "wyjasnienie_permutacji_RSA_fragment_kodu.docx"


CODE = """def run_with_permutation(run_id: int, seed: int):
    \"\"\"Shuffle all possible dimers once, then scan the permutation.\"\"\"
    rng = random.Random(seed)
    order = list(range(len(EDGES)))
    rng.shuffle(order)
    occ = [False] * (L * L)
    dsu = DSU(L * L)
    orient_grid = np.zeros((L, L), dtype=np.uint8)

    dimers = h_count = v_count = 0
    cp = math.nan
    dimers_at_cp = 0
    lr_at_cp = tb_at_cp = False

    for edge_id in order:
        a, b, orient = EDGES[edge_id]
        if occ[a] or occ[b]:
            continue

        dimers += 1
        if orient == "H":
            h_count += 1
        else:
            v_count += 1

        lr, tb = add_dimer(occ, dsu, a, b, L)
        orient_grid.flat[a] = 1 if orient == "H" else 2
        orient_grid.flat[b] = 1 if orient == "H" else 2

        if math.isnan(cp) and (lr or tb):
            cp = 2 * dimers / (L * L)
            dimers_at_cp = dimers
            lr_at_cp, tb_at_cp = lr, tb
            percolation_grid = orient_grid.copy()

    cj = 2 * dimers / (L * L)"""


def set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D9E2EC")
        borders.append(tag)
    table._tbl.tblPr.append(borders)


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = "Wyjaśnienie permutacji w symulacji RSA dimerów"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if footer.runs:
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(100, 116, 139)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run("Na czym polega permutacja w symulacji RSA dimerów")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph(
        "Osobne objaśnienie: intuicja, fragment kodu Pythona i opis działania krok po kroku."
    )
    p.runs[0].font.color.rgb = RGBColor(80, 90, 105)


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_shading(cell, "F7F9FC")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(20, 30, 40)


def add_two_col_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    set_table_width(table, [2.0, 4.5])
    headers = table.rows[0].cells
    headers[0].text = "Element kodu"
    headers[1].text = "Co robi"
    for cell in headers:
        set_shading(cell, "F2F4F7")
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def main() -> None:
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("Najkrócej: co to jest permutacja?", level=1)
    doc.add_paragraph(
        "Permutacja to losowe przestawienie kolejności elementów na liście. "
        "W tej symulacji elementem listy jest jedno możliwe położenie dimeru: dwa sąsiednie punkty kratowe plus informacja, czy dimer jest poziomy, czy pionowy."
    )
    doc.add_paragraph(
        "Dla kratki 100x100 z twardą granicą lista kandydatów ma 19 800 możliwych dimerów: "
        "9 900 poziomych i 9 900 pionowych. Permutacja oznacza: weź te 19 800 kandydatów, potasuj ich kolejność losowo i sprawdzaj po kolei."
    )

    doc.add_heading("Prosty przykład na małej liście", level=1)
    doc.add_paragraph(
        "Załóżmy, że mamy tylko pięć możliwych kandydatów: A, B, C, D, E. "
        "Bez permutacji można za każdym razem losować kolejnego kandydata. Z permutacją robimy najpierw jedną losową kolejność, na przykład:"
    )
    p = doc.add_paragraph()
    run = p.add_run("D, A, E, B, C")
    run.font.name = "Courier New"
    run.font.size = Pt(11)
    run.bold = True
    doc.add_paragraph(
        "Następnie idziemy po tej kolejności. Jeśli kandydat pasuje do wolnych punktów kratowych, dimer zostaje położony. "
        "Jeśli koliduje z wcześniej położonym dimerem, jest pomijany. Ważne: każdy kandydat jest sprawdzany najwyżej raz."
    )

    doc.add_heading("Jak to działa w RSA dimerów", level=1)
    for text in (
        "Najpierw program zna pełną listę możliwych dimerów. W kodzie ta lista nazywa się EDGES.",
        "Każdy element EDGES zawiera dwa indeksy punktów kratowych, które dimer miałby zająć, oraz orientację: H dla poziomego albo V dla pionowego.",
        "Program tworzy listę numerów od 0 do len(EDGES)-1. To są identyfikatory wszystkich możliwych dimerów.",
        "Następnie tasuje tę listę funkcją rng.shuffle(order). To właśnie jest moment wykonania permutacji.",
        "Po przetasowaniu program przechodzi po order. Każdy edge_id wskazuje jeden konkretny możliwy dimer.",
        "Jeśli oba punkty są wolne, dimer jest akceptowany. Jeśli chociaż jeden punkt jest zajęty, kandydat jest odrzucany i program idzie dalej.",
    ):
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Fragment kodu Pythona, który uruchamia permutację", level=1)
    add_code_block(doc, CODE)

    doc.add_heading("Dokładny opis fragmentu kodu", level=1)
    add_two_col_table(
        doc,
        [
            ("rng = random.Random(seed)", "Tworzy generator liczb losowych. Seed pozwala powtórzyć dokładnie tę samą symulację."),
            ("order = list(range(len(EDGES)))", "Tworzy listę numerów wszystkich możliwych położeń dimeru. To jeszcze nie jest permutacja, tylko lista w kolejności 0, 1, 2, 3..."),
            ("rng.shuffle(order)", "To jest właściwa permutacja: losowe przetasowanie kolejności kandydatów."),
            ("for edge_id in order:", "Program przechodzi po przetasowanej liście. Kolejność sprawdzania dimerów jest losowa, ale ustalona dla danego przebiegu."),
            ("a, b, orient = EDGES[edge_id]", "Pobiera konkretny dimer: pierwszy punkt a, drugi punkt b oraz orientację H albo V."),
            ("if occ[a] or occ[b]: continue", "Jeśli którykolwiek z dwóch punktów jest już zajęty, dimer nie może być położony i program przechodzi do następnego kandydata."),
            ("dimers += 1", "Zlicza zaakceptowany dimer. Ta linia wykona się tylko wtedy, gdy oba punkty były wolne."),
            ("h_count / v_count", "Zlicza, ile zaakceptowano dimerów poziomych i pionowych."),
            ("add_dimer(...)", "Zajmuje punkty na siatce i aktualizuje strukturę klastrów potrzebną do wykrycia perkolacji."),
            ("if math.isnan(cp) and (lr or tb):", "Sprawdza, czy perkolacja pojawiła się po raz pierwszy. lr oznacza lewa-prawa, tb oznacza góra-dół."),
            ("cp = 2 * dimers / (L * L)", "Zapisuje próg perkolacji jako pokrycie powierzchni w momencie pierwszego połączenia brzegów."),
            ("cj = 2 * dimers / (L * L)", "Po przejściu całej permutacji zapisuje końcowe pokrycie, czyli jamming threshold."),
        ],
    )

    doc.add_heading("Dlaczego to daje stan jammed?", level=1)
    doc.add_paragraph(
        "Bo lista EDGES zawiera wszystkie dopuszczalne położenia dimeru przy hard boundary condition. "
        "Jeśli po przejściu całej przetasowanej listy jakiś dimer nie został położony, to znaczy, że w momencie jego sprawdzania kolidował z już zajętymi punktami. "
        "Na końcu nie zostaje żaden legalny kandydat, którego program nie rozważył."
    )
    doc.add_paragraph(
        "Właśnie dlatego wersja z permutacją jest wygodna: nie trzeba zgadywać, po ilu nieudanych próbach uznać, że układ jest zablokowany. "
        "Koniec listy jest jednoznacznym kryterium zakończenia."
    )

    doc.add_heading("Jednozdaniowa intuicja", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "Permutacja w tej symulacji to po prostu losowa kolejka wszystkich możliwych dimerów: każdy dimer dostaje swoje miejsce w kolejce, a program sprawdza je po kolei."
    )
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("1F4D78")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
