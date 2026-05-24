#!/usr/bin/env python3
"""Build DOCX reports from RSA simulation outputs."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data"
FIGURES = ROOT / "figures"
REPORTS = ROOT / "reports"
PROMPT_PATH = Path(
    "/Users/piotrlunkiewicz/Desktop/Doktorat UAM/Zadania/Permutacja sprawdzamy/###Symulacja permutacja czy coś zmienia?"
)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D9E2EC")
        borders.append(tag)
    tbl_pr.append(borders)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = "RSA dimery 100x100, hard boundary, 500 przebiegow na metode"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(100, 116, 139)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(11, 37, 69)
    r.bold = True
    s = doc.add_paragraph(subtitle)
    s.runs[0].font.size = Pt(11)
    s.runs[0].font.color.rgb = RGBColor(80, 90, 105)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.1)
    table.columns[1].width = Inches(4.4)
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "Parametr"
    hdr[1].text = "Wartosc"
    for cell in hdr:
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].bold = True
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_summary_table(doc: Document, summary: list[dict[str, str]]) -> None:
    rows = [r for r in summary if r["metric"] in ("cp", "cj")]
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    headers = ["Metoda", "Wielkosc", "n", "srednia", "SD", "SEM", "95% CI", "min-max"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        set_cell_shading(cell, "F2F4F7")
        cell.paragraphs[0].runs[0].bold = True
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = r["method"].replace("_", " ")
        cells[1].text = r["metric"].upper()
        cells[2].text = r["n"]
        cells[3].text = f"{float(r['mean']):.6f}"
        cells[4].text = f"{float(r['sd']):.6f}"
        cells[5].text = f"{float(r['sem']):.6f}"
        cells[6].text = f"{float(r['ci95_low']):.6f} - {float(r['ci95_high']):.6f}"
        cells[7].text = f"{float(r['min']):.4f} - {float(r['max']):.4f}"
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_comparison_table(doc: Document, comparison: list[dict[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    headers = ["Wielkosc", "bez permutacji", "z permutacja", "roznica", "t Welcha", "p approx."]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        set_cell_shading(cell, "F2F4F7")
        cell.paragraphs[0].runs[0].bold = True
    for r in comparison:
        cells = table.add_row().cells
        cells[0].text = r["metric"].upper()
        cells[1].text = f"{float(r['mean_bez_permutacji']):.6f}"
        cells[2].text = f"{float(r['mean_z_permutacja']):.6f}"
        cells[3].text = f"{float(r['difference']):+.6f}"
        cells[4].text = f"{float(r['welch_t']):.3f}"
        cells[5].text = f"{float(r['p_approx_normal']):.3f}"


def add_picture(doc: Document, filename: str, caption: str, width: float = 6.2) -> None:
    doc.add_picture(str(FIGURES / filename), width=Inches(width))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(80, 90, 105)


def build_results_report() -> Path:
    summary = read_csv(DATA / "rsa_summary_statistics.csv")
    comparison = read_csv(DATA / "rsa_method_comparison.csv")
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "Porownanie RSA dimerow: bez permutacji i z permutacja",
        "Komorka Monte Carlo 100x100, hard boundary condition, po 500 pelnych przebiegow dla kazdej metody.",
    )

    doc.add_heading("Zakres i zrodla", level=1)
    doc.add_paragraph(
        "Zastosowano lokalny prompt z katalogu 'Permutacja sprawdzamy' oraz plik zadania Prof. Andrzeja Sikorskiego z 18.02.2026. "
        "Repozytorium GitHub wskazane w poleceniu istnieje, ale w chwili pracy nie zawieralo plikow do odczytania; dlatego prompt zostal pobrany z lokalnego pliku informacyjnego."
    )
    add_kv_table(
        doc,
        [
            ("L", "100"),
            ("Liczba wezlow", "10 000"),
            ("Metody", "bez wstepnej permutacji; z losowa permutacja pelnej listy kandydatow"),
            ("Liczba przebiegow", "500 na metode"),
            ("Warunki brzegowe", "hard boundary, bez zawijania przez krawedzie"),
            ("Perkolacja Cp", "pierwsze pokrycie, przy ktorym klaster zajetych wezlow laczy lewa-prawa albo gora-dol"),
            ("Jamming Cj", "koncowe pokrycie, gdy nie istnieje juz zaden legalny dimer"),
        ],
    )

    doc.add_heading("Wyniki liczbowe", level=1)
    add_summary_table(doc, summary)
    doc.add_paragraph(
        "Srednie Cj sa praktycznie identyczne: 0.906565 bez permutacji oraz 0.906510 z permutacja. "
        "Dla Cp widac nieco wieksza losowa fluktuacje, ale roznica srednich wynosi tylko -0.001567 na korzysc serii z permutacja."
    )
    add_comparison_table(doc, comparison)

    doc.add_heading("Wykresy i konfiguracje", level=1)
    add_picture(doc, "boxplot_cp_cj.png", "Rozklady Cp i Cj dla dwoch metod.")
    add_picture(doc, "hist_cp.png", "Histogram progu perkolacji Cp.")
    add_picture(doc, "hist_cj.png", "Histogram progu jammingu Cj.")
    add_picture(doc, "ratios_prompt_metrics.png", "Wskazniki wymagane w promptcie: Cp/dp, Cj/dj, Cp/L, Cj/L, Cp/Cj.")
    add_picture(doc, "konfiguracja_perkolacja_bez_permutacji.png", "Przykladowa konfiguracja przy pierwszej perkolacji: bez permutacji.", 5.4)
    add_picture(doc, "konfiguracja_jamming_bez_permutacji.png", "Przykladowa konfiguracja w stanie jammed: bez permutacji.", 5.4)
    add_picture(doc, "konfiguracja_perkolacja_z_permutacja.png", "Przykladowa konfiguracja przy pierwszej perkolacji: z permutacja.", 5.4)
    add_picture(doc, "konfiguracja_jamming_z_permutacja.png", "Przykladowa konfiguracja w stanie jammed: z permutacja.", 5.4)

    doc.add_heading("Pliki wynikowe", level=1)
    for item in (
        "data/rsa_500_runs_per_method.csv",
        "data/rsa_summary_statistics.csv",
        "data/rsa_method_comparison.csv",
        "logs/dziennik_symulacji.md",
        "scripts/rsa_permutation_comparison.py",
    ):
        doc.add_paragraph(item, style="List Bullet")

    out = REPORTS / "raport_wyniki_RSA_dimery_permutacja_vs_bez.docx"
    doc.save(out)
    return out


def build_discussion_report() -> Path:
    comparison = read_csv(DATA / "rsa_method_comparison.csv")
    cp = next(r for r in comparison if r["metric"] == "cp")
    cj = next(r for r in comparison if r["metric"] == "cj")

    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "Permutacja w RSA dimerow: wyjasnienie i dyskusja",
        "Osobne opracowanie o tym, co oznacza permutacja w tej symulacji i dlaczego wyniki moga byc takie same albo rozne.",
    )

    doc.add_heading("Na czym polega permutacja", level=1)
    doc.add_paragraph(
        "W tej symulacji kandydatem jest konkretne polozenie dimeru: para sasiadujacych punktow kratowych oraz orientacja pozioma albo pionowa. "
        "Dla siatki 100x100 i twardych brzegow istnieje 2*L*(L-1), czyli 19 800 takich kandydatow. "
        "Permutacja oznacza jednorazowe ustawienie tych 19 800 kandydatow w losowej kolejnosci."
    )
    doc.add_paragraph(
        "Algorytm z permutacja przechodzi po tej liscie tylko raz. Jesli oba punkty dimeru sa wolne, dimer zostaje przyjety; jesli choc jeden punkt jest juz zajety, kandydat jest pomijany. "
        "Koniec listy daje stan zablokowany, bo kazde mozliwe polozenie zostalo juz sprawdzone."
    )

    doc.add_heading("Czym rozni sie wersja bez permutacji", level=1)
    doc.add_paragraph(
        "W wersji bez wstepnej permutacji nie tworzy sie jednej pelnej kolejnosci na starcie. Kolejny dimer jest losowany z aktualnie legalnych kandydatow. "
        "Jest to odpowiednik klasycznego RSA z odrzucaniem po odfiltrowaniu prob, ktore i tak nic nie zmieniaja w konfiguracji."
    )
    doc.add_paragraph(
        "Odrzucone proby zmieniaja liczbe prob potrzebnych do dojscia do konca, ale nie zmieniaja pokrycia po kolejnych akceptacjach. "
        "Dlatego dla Cj i Cp liczonych jako pokrycie powierzchni, a nie jako czas liczony w probach, obie procedury powinny dawac ten sam rozklad statystyczny."
    )

    doc.add_heading("Wplyw na wynik", level=1)
    doc.add_paragraph(
        f"Dla 500 przebiegow na metode srednia roznica Cp wyniosla {float(cp['difference']):+.6f}, "
        f"a przyblizone p dla testu Welcha wynioslo {float(cp['p_approx_normal']):.3f}. "
        f"Dla Cj srednia roznica wyniosla {float(cj['difference']):+.6f}, "
        f"z p okolo {float(cj['p_approx_normal']):.3f}."
    )
    doc.add_paragraph(
        "Te liczby nie pokazuja istotnej roznicy miedzy metodami. Najrozsadniejsza interpretacja jest taka, ze obserwowane odchylenia sa zwyklym szumem Monte Carlo. "
        "Cj jest szczegolnie stabilne, bo jamming na duzej siatce ma waski rozklad; Cp ma szerszy rozklad, bo zalezy od momentu powstania pierwszego klastra spinajacego brzegi."
    )

    doc.add_heading("Kiedy wyniki moglyby sie roznic", level=1)
    doc.add_paragraph(
        "Roznice pojawilyby sie, gdyby algorytm z permutacja nie byl rownowazny procesowi RSA, na przyklad gdyby kandydaci nie mieli jednakowych szans, "
        "gdyby lista nie obejmowala wszystkich polozen, gdyby po zajeciu punktow usuwano zbyt wiele albo zbyt malo kandydatow, albo gdyby klasyczna wersja konczyla sie po arbitralnej liczbie porazek zamiast po rzeczywistym braku wolnego dimeru."
    )
    doc.add_paragraph(
        "W praktyce permutacja jest wiec przede wszystkim przyspieszeniem i sposobem unikniecia problemu: ile odrzuconych prob wystarczy, by uznac stan za jammed. "
        "Nie powinna zmieniac Cj ani Cp jako wielkosci geometrycznych, o ile zachowana jest ta sama przestrzen kandydatow i to samo kryterium akceptacji."
    )

    doc.add_heading("Co jeszcze warto zrobic", level=1)
    for item in (
        "uruchomic parowane przebiegi na tej samej losowej kolejnosci kandydatow, aby pokazac rownowaznosc algorytmiczna krok po kroku",
        "powtorzyc badanie dla kilku rozmiarow L, np. 50, 100, 150, 200, i zrobic ekstrapolacje rozmiarowa",
        "porownac nie tylko Cp i Cj, ale tez liczbe dimerow poziomych/pionowych oraz parametr porzadku orientacyjnego",
        "dodac rozklad rozmiarow pustych klastrow w stanie jammed",
        "osobno analizowac czas w probach dla klasycznego RSA z odrzucaniem, bo tu permutacja rzeczywiscie zmienia koszt obliczeniowy",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Uwaga o repozytorium GitHub", level=1)
    doc.add_paragraph(
        "Wskazane repozytorium git@github.com:piotr-collab/Sym-z-permutacj-i-bez---sprawdzenie.git zostalo sprawdzone. "
        "Metadane GitHuba wskazywaly rozmiar 0 i brak plikow do odczytania, dlatego lokalny plik informacyjny zostal potraktowany jako tresc promptu."
    )

    out = REPORTS / "opracowanie_permutacja_w_RSA_wyjasnienie_i_dyskusja.docx"
    doc.save(out)
    return out


def main() -> None:
    REPORTS.mkdir(exist_ok=True)
    results = [build_results_report(), build_discussion_report()]
    if PROMPT_PATH.exists():
        (REPORTS / "prompt_lokalny.txt").write_text(PROMPT_PATH.read_text(encoding="utf-8"), encoding="utf-8")
    for path in results:
        print(path)


if __name__ == "__main__":
    main()
